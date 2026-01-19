"""
Agentic Orchestrator - True Multi-Agent Pipeline Coordination

Key improvements over v1:
- Real progress tracking with stage-by-stage callbacks
- Per-agent timeout enforcement
- Graceful degradation for optional agents
- Error isolation - one agent failure doesn't crash everything
- Checkpoint/resume capability (future)
- Parallel execution support for independent agents (future)
"""
import os
import pandas as pd
from pathlib import Path
from typing import Dict, Any, Optional, List, Callable
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

from agents.base import (
    AgentBase, AgentResult, AgentConfig, AgentStatus,
    PipelineContext, ProgressCallback
)


class PipelineStatus(Enum):
    """Overall pipeline status."""
    NOT_STARTED = "not_started"
    RUNNING = "running"
    SUCCESS = "success"
    PARTIAL_SUCCESS = "partial_success"  # Some optional agents failed
    FAILED = "failed"
    CANCELLED = "cancelled"


@dataclass
class StageInfo:
    """Information about a pipeline stage."""
    name: str
    agent: AgentBase
    required: bool = True
    depends_on: List[str] = field(default_factory=list)


@dataclass
class PipelineResult:
    """Complete result from the pipeline."""
    success: bool
    status: PipelineStatus
    harmonized_df: Optional[pd.DataFrame] = None
    qc_report: Optional[pd.DataFrame] = None
    transformation_summary: list = field(default_factory=list)
    review_result: Optional[Dict[str, Any]] = None
    run_metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None
    error_stage: Optional[str] = None
    stage_results: Dict[str, AgentResult] = field(default_factory=dict)


# Callback type for overall progress: (current_stage, total_stages, stage_name, stage_status, message, overall_pct)
PipelineProgressCallback = Callable[[int, int, str, AgentStatus, str, float], None]


class AgenticOrchestrator:
    """
    Orchestrator for the Concordia harmonization pipeline.

    Coordinates agents with:
    - Real-time progress tracking
    - Timeout enforcement per agent
    - Graceful handling of optional agent failures
    - Clear error attribution
    """

    def __init__(
        self,
        output_dir: Optional[str] = None,
        use_llm: bool = True,
        progress_callback: Optional[PipelineProgressCallback] = None
    ):
        """
        Initialize the orchestrator.

        Args:
            output_dir: Directory for output files
            use_llm: Whether to enable LLM features
            progress_callback: Callback for progress updates
        """
        self.output_dir = Path(output_dir) if output_dir else Path("output")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.use_llm = use_llm
        self.progress_callback = progress_callback
        self._status = PipelineStatus.NOT_STARTED
        self._stages: List[StageInfo] = []
        self._current_stage_idx = 0

    def _notify_progress(
        self,
        stage_name: str,
        stage_status: AgentStatus,
        message: str,
        stage_progress: float
    ):
        """Calculate and emit overall progress."""
        if not self.progress_callback:
            return

        total_stages = len(self._stages)
        current_idx = self._current_stage_idx

        # Overall progress: completed stages + current stage progress
        completed_weight = current_idx / total_stages if total_stages > 0 else 0
        current_weight = (stage_progress / total_stages) if total_stages > 0 else 0
        overall_pct = completed_weight + current_weight

        self.progress_callback(
            current_idx + 1,
            total_stages,
            stage_name,
            stage_status,
            message,
            overall_pct
        )

    def _create_agent_callback(self, stage_name: str) -> ProgressCallback:
        """Create a progress callback bound to a specific stage."""
        def callback(agent_name: str, status: AgentStatus, message: str, progress: float):
            self._notify_progress(stage_name, status, message, progress)
        return callback

    def _build_stages(self, enable_llm_review: bool) -> List[StageInfo]:
        """Build the list of pipeline stages with configured agents."""
        from agents.ingest_agent import IngestAgent
        from agents.map_agent import MapAgent
        from agents.harmonize_agent import HarmonizeAgent
        from agents.qc_agent import QCAgent
        from agents.review_agent import ReviewAgent

        stages = []

        # Stage 1: Ingest
        ingest_config = AgentConfig(
            timeout_seconds=60.0,
            required=True
        )
        stages.append(StageInfo(
            name="ingest",
            agent=IngestAgent(
                config=ingest_config,
                progress_callback=self._create_agent_callback("ingest")
            ),
            required=True
        ))

        # Stage 2: Map
        map_config = AgentConfig(
            timeout_seconds=60.0,
            required=True
        )
        stages.append(StageInfo(
            name="map",
            agent=MapAgent(
                config=map_config,
                progress_callback=self._create_agent_callback("map")
            ),
            required=True,
            depends_on=["ingest"]
        ))

        # Stage 3: Harmonize
        harmonize_config = AgentConfig(
            timeout_seconds=120.0,  # Longer for potential LLM calls
            max_retries=1,  # Retry once on transient failures
            required=True
        )
        stages.append(StageInfo(
            name="harmonize",
            agent=HarmonizeAgent(
                config=harmonize_config,
                use_llm_fallback=self.use_llm,
                progress_callback=self._create_agent_callback("harmonize")
            ),
            required=True,
            depends_on=["map"]
        ))

        # Stage 4: QC
        qc_config = AgentConfig(
            timeout_seconds=60.0,
            required=True
        )
        stages.append(StageInfo(
            name="qc",
            agent=QCAgent(
                config=qc_config,
                progress_callback=self._create_agent_callback("qc")
            ),
            required=True,
            depends_on=["harmonize"]
        ))

        # Stage 5: LLM Review (optional)
        if self.use_llm and enable_llm_review:
            review_config = AgentConfig(
                timeout_seconds=90.0,  # LLM calls can be slow
                max_retries=1,
                required=False  # Optional - failure doesn't stop pipeline
            )
            stages.append(StageInfo(
                name="review",
                agent=ReviewAgent(
                    config=review_config,
                    progress_callback=self._create_agent_callback("review")
                ),
                required=False,  # Graceful degradation
                depends_on=["qc"]
            ))

        return stages

    def run(
        self,
        data_file: str,
        dictionary_file: Optional[str] = None,
        save_outputs: bool = True,
        enable_llm_review: bool = True
    ) -> PipelineResult:
        """
        Run the complete harmonization pipeline.

        Args:
            data_file: Path to source data file
            dictionary_file: Optional path to data dictionary
            save_outputs: Whether to save output files
            enable_llm_review: Whether to run LLM review stage

        Returns:
            PipelineResult with all outputs and metadata
        """
        run_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        self._status = PipelineStatus.RUNNING

        run_metadata = {
            "run_id": run_id,
            "start_time": datetime.now().isoformat(),
            "data_file": data_file,
            "dictionary_file": dictionary_file,
            "stages": {},
            "llm_enabled": self.use_llm
        }

        # Initialize context with input files
        context = PipelineContext({
            "data_file": data_file,
            "dictionary_file": dictionary_file,
            "run_id": run_id,
            "output_dir": str(self.output_dir)
        })

        # Build pipeline stages
        self._stages = self._build_stages(enable_llm_review)
        stage_results: Dict[str, AgentResult] = {}
        failed_required = False
        failed_optional = False

        # Execute stages sequentially
        for idx, stage_info in enumerate(self._stages):
            self._current_stage_idx = idx
            stage_name = stage_info.name
            agent = stage_info.agent

            # Log stage start
            total = len(self._stages)
            print(f"[{idx + 1}/{total}] Running {agent.name}...")

            # Check dependencies
            deps_met = all(
                stage_results.get(dep, AgentResult(success=False)).success
                for dep in stage_info.depends_on
            )

            if not deps_met:
                print(f"    ⚠ Skipping {stage_name}: dependencies not met")
                result = AgentResult(
                    success=False,
                    error="Dependencies not met",
                    error_type="DependencyError"
                )
                stage_results[stage_name] = result
                context.set_stage_result(stage_name, result)

                if stage_info.required:
                    failed_required = True
                    break
                continue

            # Run the agent
            try:
                result = agent.run(context)
                stage_results[stage_name] = result
                context.set_stage_result(stage_name, result)

                # Store stage metadata
                run_metadata["stages"][stage_name] = {
                    "success": result.success,
                    "execution_time_ms": result.execution_time_ms,
                    "retries_used": result.retries_used,
                    **result.metadata
                }

                if result.success:
                    # Update context with agent outputs
                    if result.data:
                        context.update(result.data)
                    print(f"    ✓ {stage_name} complete ({result.execution_time_ms}ms)")
                else:
                    error_msg = result.error or "Unknown error"
                    if stage_info.required:
                        print(f"    ✗ {stage_name} failed: {error_msg}")
                        failed_required = True
                        break
                    else:
                        print(f"    ⚠ {stage_name} failed (optional): {error_msg}")
                        failed_optional = True

            except Exception as e:
                import traceback
                error_msg = f"{type(e).__name__}: {str(e)}"
                result = AgentResult(
                    success=False,
                    error=error_msg,
                    error_type=type(e).__name__,
                    metadata={"traceback": traceback.format_exc()}
                )
                stage_results[stage_name] = result
                context.set_stage_result(stage_name, result)

                run_metadata["stages"][stage_name] = {
                    "success": False,
                    "error": error_msg
                }

                if stage_info.required:
                    print(f"    ✗ {stage_name} crashed: {error_msg}")
                    failed_required = True
                    break
                else:
                    print(f"    ⚠ {stage_name} crashed (optional): {error_msg}")
                    failed_optional = True

        # Determine final status
        run_metadata["end_time"] = datetime.now().isoformat()

        if failed_required:
            self._status = PipelineStatus.FAILED
            failed_stage = next(
                (name for name, r in stage_results.items() if not r.success),
                "Unknown"
            )
            run_metadata["success"] = False

            return PipelineResult(
                success=False,
                status=self._status,
                error=stage_results.get(failed_stage, AgentResult(success=False)).error,
                error_stage=failed_stage,
                run_metadata=run_metadata,
                stage_results=stage_results
            )

        # Pipeline succeeded (possibly with optional failures)
        if failed_optional:
            self._status = PipelineStatus.PARTIAL_SUCCESS
        else:
            self._status = PipelineStatus.SUCCESS

        run_metadata["success"] = True

        # Extract results from context
        harmonized_df = context.get("harmonized_df")
        qc_report = context.get("qc_report")
        transformation_summary = context.get("transformation_summary", [])
        review_result = context.get("review_result")

        # Save outputs if requested
        output_files = {}
        if save_outputs and harmonized_df is not None:
            output_files = self._save_outputs(
                run_id,
                harmonized_df,
                qc_report,
                transformation_summary,
                run_metadata
            )
            run_metadata["output_files"] = output_files
            print(f"\n✓ Pipeline complete! Outputs saved to: {self.output_dir}")

        return PipelineResult(
            success=True,
            status=self._status,
            harmonized_df=harmonized_df,
            qc_report=qc_report,
            transformation_summary=transformation_summary,
            review_result=review_result,
            run_metadata=run_metadata,
            stage_results=stage_results
        )

    def _save_outputs(
        self,
        run_id: str,
        harmonized_df: pd.DataFrame,
        qc_report: Optional[pd.DataFrame],
        transformation_summary: list,
        run_metadata: Dict
    ) -> Dict[str, str]:
        """Save all output files."""
        output_files = {}
        trial_id = run_metadata.get('stages', {}).get('ingest', {}).get('trial_id', 'UNKNOWN')

        # 1. Harmonized dataset (CSV)
        harmonized_path = self.output_dir / f"{trial_id}_DM_harmonized_{run_id}.csv"
        harmonized_df.to_csv(harmonized_path, index=False)
        output_files["harmonized_csv"] = str(harmonized_path)

        # 2. QC Report (CSV)
        if qc_report is not None:
            qc_path = self.output_dir / f"{trial_id}_QC_report_{run_id}.csv"
            qc_report.to_csv(qc_path, index=False)
            output_files["qc_csv"] = str(qc_path)

        # 3. Transformation Summary (CSV)
        if transformation_summary:
            transform_path = self.output_dir / f"{trial_id}_transformation_summary_{run_id}.csv"
            pd.DataFrame(transformation_summary).to_csv(transform_path, index=False)
            output_files["transformation_csv"] = str(transform_path)

        # 4. DOCX Report
        docx_path = self.output_dir / f"{trial_id}_Harmonization_Report_{run_id}.docx"
        self._generate_docx_report(
            docx_path,
            harmonized_df,
            qc_report,
            transformation_summary,
            run_metadata
        )
        output_files["report_docx"] = str(docx_path)

        return output_files

    def _generate_docx_report(
        self,
        output_path: Path,
        harmonized_df: pd.DataFrame,
        qc_report: Optional[pd.DataFrame],
        transformation_summary: list,
        run_metadata: Dict
    ):
        """Generate the Harmonization Transformation Report as DOCX."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading('Harmonization Transformation Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            ingest_meta = run_metadata.get('stages', {}).get('ingest', {})
            harmonize_meta = run_metadata.get('stages', {}).get('harmonize', {})

            trial_id = ingest_meta.get('trial_id', 'Unknown')
            source_filename = ingest_meta.get('source_filename', 'Unknown')
            rows_in = ingest_meta.get('rows', 0)
            rows_out = harmonize_meta.get('rows_out', 0)
            rows_dropped = rows_in - rows_out if rows_in and rows_out else 0

            doc.add_paragraph(f"Trial: {trial_id}")
            doc.add_paragraph(f"Input: {source_filename}")
            doc.add_paragraph(f"Rows in input: {rows_in} | Rows in output: {rows_out} | Rows dropped: {rows_dropped}")

            # Execution info
            doc.add_paragraph(f"Pipeline version: v2 (Agentic Architecture)")
            doc.add_paragraph(f"Run ID: {run_metadata.get('run_id', 'Unknown')}")

            # Stage execution times
            doc.add_heading('Execution Summary', level=1)
            for stage_name, stage_info in run_metadata.get('stages', {}).items():
                status = "✓" if stage_info.get('success') else "✗"
                time_ms = stage_info.get('execution_time_ms', 0)
                doc.add_paragraph(f"{status} {stage_name}: {time_ms}ms")

            # Dictionary info
            if ingest_meta.get('dictionary_loaded'):
                dict_filename = ingest_meta.get('dictionary_filename', 'Unknown')
                doc.add_paragraph(f"Dictionary used: {dict_filename}")
            else:
                doc.add_paragraph("Dictionary used: None")

            doc.add_paragraph()

            # Output Schema Section
            doc.add_heading('1. Output Schema', level=1)
            from config.schema import OUTPUT_SCHEMA
            doc.add_paragraph(', '.join(OUTPUT_SCHEMA))

            # Variable Transformation Table
            doc.add_heading('2. Variable-Level Transformations', level=1)

            if transformation_summary:
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'
                headers = ['Variable', 'Source', 'Operation', 'Details', 'Changed', '%', 'Missing']
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header

                for item in transformation_summary:
                    row = table.add_row()
                    row.cells[0].text = str(item.get('variable', ''))
                    source_col = item.get('source_column', '') or ''
                    row.cells[1].text = str(source_col) if source_col else '(derived)'
                    row.cells[2].text = str(item.get('mapping_operation', 'Copy'))
                    row.cells[3].text = str(item.get('transform_operation', ''))[:50]
                    row.cells[4].text = str(item.get('rows_changed', 0))
                    row.cells[5].text = f"{item.get('percent_changed', 0):.1f}%"
                    row.cells[6].text = str(item.get('missing_count', 0))

            # QC Report Section
            doc.add_heading('3. QC Report', level=1)

            # Check if there are any QC issues - use the DataFrame directly as source of truth
            # (metadata key is 'total_qc_issues' from QC agent, but DataFrame is authoritative)
            has_qc_issues = qc_report is not None and len(qc_report) > 0

            if has_qc_issues:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                headers = ['TRIAL', 'Issue Type', 'Variable', 'Rows Affected', 'Notes']
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header

                for _, issue in qc_report.iterrows():
                    row = table.add_row()
                    row.cells[0].text = str(trial_id)
                    row.cells[1].text = str(issue.get('issue_type', ''))
                    row.cells[2].text = str(issue.get('variable', ''))
                    row.cells[3].text = str(issue.get('n_rows_affected', 0))
                    row.cells[4].text = str(issue.get('notes', ''))[:100]
            else:
                doc.add_paragraph("No QC issues found.")

            # Files Produced
            doc.add_heading('4. Files Produced', level=1)
            doc.add_paragraph(f"Harmonized output: {trial_id}_DM_harmonized_*.csv")
            doc.add_paragraph(f"QC report: {trial_id}_QC_report_*.csv")
            doc.add_paragraph(f"Transformation report: This document")

            doc.save(output_path)

        except ImportError:
            # Fallback to text
            txt_path = output_path.with_suffix('.txt')
            with open(txt_path, 'w') as f:
                f.write("HARMONIZATION REPORT\n")
                f.write(f"Trial: {run_metadata.get('stages', {}).get('ingest', {}).get('trial_id', 'Unknown')}\n")


def run_pipeline(
    data_file: str,
    dictionary_file: Optional[str] = None,
    output_dir: Optional[str] = None,
    progress_callback: Optional[PipelineProgressCallback] = None
) -> PipelineResult:
    """
    Convenience function to run the complete pipeline.
    """
    orchestrator = AgenticOrchestrator(
        output_dir=output_dir,
        progress_callback=progress_callback
    )
    return orchestrator.run(data_file, dictionary_file)
